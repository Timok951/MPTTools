"""
Сбор полного листинга исходного кода проекта TIP (MPT Tools) в DOCX
с таблицей «Таблица 1 – Описание модулей» (№, файл, язык, строки, КБ, описание).

Запуск из корня репозитория:
  .venv\\Scripts\\python.exe allcode.py
  .venv\\Scripts\\python.exe allcode.py -o MPT_Tools_full_code.docx
"""

from __future__ import annotations

import argparse
import ast
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SOURCE_ROOT = SCRIPT_DIR / "TIP"
DEFAULT_OUTPUT = SCRIPT_DIR / "full_code_report.docx"

EXCLUDE_DIRS = {
    "__pycache__",
    ".git",
    ".idea",
    ".vscode",
    ".cursor",
    "venv",
    ".venv",
    "env",
    "migrations",
    "node_modules",
    "mcps",
    "site-packages",
}

INCLUDE_EXTENSIONS = (".py", ".html", ".css", ".js", ".ts", ".tsx", ".json")

# Дополнительно: не тащить мусор
EXCLUDE_FILE_NAMES = {".env", ".env.example"}
EXCLUDE_PATH_PARTS = {".git", ".venv", "node_modules"}


def lang_for_suffix(path: Path) -> str:
    s = path.suffix.lower()
    if s == ".py":
        return "Python"
    if s == ".html":
        return "HTML"
    if s == ".css":
        return "CSS"
    if s == ".js":
        return "JavaScript"
    if s == ".ts":
        return "TypeScript"
    if s == ".tsx":
        return "TypeScript JSX"
    if s == ".json":
        return "JSON"
    return s.lstrip(".") or "—"


def _clean_docstring(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\s+", " ", s)
    if len(s) > 400:
        s = s[:397] + "…"
    return s or "—"


def module_docstring_py(src: str) -> str | None:
    try:
        tree = ast.parse(src)
    except SyntaxError:
        return None
    doc = ast.get_docstring(tree, clean=False)
    if doc:
        return _clean_docstring(doc)
    return None


def describe_file(rel_posix: str, filename: str, lang: str) -> str:
    """Краткое назначение модуля, если нет осмысленного docstring."""
    lower = rel_posix.lower()
    base = filename.lower()

    rules: list[tuple[str, str]] = [
        ("manage.py", "Точка входа Django: сервер, миграции, управляющие команды."),
        ("wsgi.py", "WSGI-конфигурация для развёртывания под совместимыми серверами."),
        ("asgi.py", "ASGI-конфигурация для асинхронных серверов."),
        ("settings.py", "Настройки проекта Django: приложения, БД, middleware, безопасность."),
        ("urls.py", "Карта URL: привязка путей к представлениям или включение других urlconf."),
        ("api_urls.py", "Маршруты REST API."),
        ("models.py", "Модели данных Django (ORM) и бизнес-сущности предметной области."),
        ("admin.py", "Регистрация моделей в административной панели Django."),
        ("views.py", "Представления: обработка HTTP-запросов, выдача HTML и перенаправления."),
        ("portal_views.py", "Представления портала администрирования объектов."),
        ("forms.py", "Формы ввода, валидация и связь с моделями."),
        ("portal_forms.py", "Формы для разделов портала."),
        ("tests.py", "Модульные и интеграционные тесты."),
        ("test_", "Тестовый модуль (проверки поведения приложения)."),
        ("middleware.py", "Промежуточная обработка запросов и ответов."),
        ("signals.py", "Сигналы Django: реакции на сохранение/удаление объектов."),
        ("serializers.py", "Сериализация и десериализация данных REST API."),
        ("viewsets.py", "ViewSet-ы REST API (CRUD и действия)."),
        ("permissions.py", "Правила доступа к API."),
        ("exceptions.py", "Исключения и форматы ошибок API."),
        ("auth_views.py", "Точки аутентификации для API."),
        ("apps.py", "Конфигурация приложения Django."),
        ("context_processors.py", "Дополнительный контекст для шаблонов."),
        ("templatetags/", "Пользовательские теги и фильтры шаблонов."),
        ("management/commands/", "Пользовательские команды manage.py."),
        ("async_tasks.py", "Вспомогательная логика фоновых или отложенных задач."),
        ("mail_out.py", "Отправка служебной электронной почты."),
        ("registration_domains.py", "Правила регистрации по доменам e-mail."),
        ("schedule_utils.py", "Утилиты для расписаний и циклов."),
        ("periodic_usage.py", "Логика периодического учёта потребления."),
        ("portal_log.py", "Журналирование действий портала."),
        ("notification_utils.py", "Формирование и доставка уведомлений пользователям."),
        ("backup_utils.py", "Резервное копирование данных."),
        ("quality_report.py", "Отчёт о качестве/целостности данных."),
        ("portal_urls.py", "Маршруты портала."),
        ("api/", "Код REST API."),
    ]

    for needle, text in rules:
        if needle.endswith("/") and needle in lower.replace("\\", "/"):
            return text
        if base == needle or lower.endswith("/" + needle):
            return text
        if needle.endswith(".py") and base == needle:
            return text
        if needle == "test_" and base.startswith("test_"):
            return text

    if lang == "HTML":
        if "templates/" in lower:
            return "HTML-шаблон Django: разметка страницы или фрагмента интерфейса."
        return "HTML-разметка."

    if lang == "CSS":
        return "Таблица стилей."

    if lang in ("JavaScript", "TypeScript", "TypeScript JSX"):
        return "Клиентский сценарий или модуль."

    if lang == "JSON":
        return "Данные в формате JSON (конфигурация или сообщения)."

    return f"Исходный файл приложения ({lang})."


def description_for_file(rel_posix: str, content: str, lang: str) -> str:
    path = Path(rel_posix)
    if path.suffix.lower() == ".py":
        doc = module_docstring_py(content)
        if doc and doc != "—":
            return doc
    return describe_file(rel_posix, path.name, lang)


def should_skip_path(rel: Path) -> bool:
    parts = set(rel.parts)
    if parts & EXCLUDE_PATH_PARTS:
        return True
    for p in rel.parts:
        if p in EXCLUDE_DIRS:
            return True
    if rel.name in EXCLUDE_FILE_NAMES:
        return True
    return False


def collect_files(source_root: Path) -> list[Path]:
    out: list[Path] = []
    for p in source_root.rglob("*"):
        if not p.is_file():
            continue
        try:
            rel = p.relative_to(source_root)
        except ValueError:
            continue
        if should_skip_path(rel):
            continue
        if p.suffix.lower() not in INCLUDE_EXTENSIONS:
            continue
        out.append(p)
    return out


def folder_label(source_root: Path, rel_file: Path) -> str:
    """Подпись группы для текста «Папка «…»»."""
    root_name = source_root.name
    parent = rel_file.parent.as_posix()
    if parent in (".", ""):
        return root_name
    return f"{root_name}/{parent}"


def sort_key(p: Path, source_root: Path) -> tuple:
    rel = p.relative_to(source_root)
    parent = str(rel.parent)
    # корень проекта первым
    root_order = 0 if parent in (".", "") else 1
    return (root_order, parent.lower(), rel.name.lower())


def format_kb(kb: float) -> str:
    return f"{kb:.2f}".replace(".", ",")


def line_count_for_text(raw: str) -> int:
    if not raw:
        return 0
    return len(raw.splitlines())


def sanitize_for_word_xml(text: str) -> str:
    """Удаляет NUL и прочие символы, недопустимые в XML (python-docx / lxml)."""
    if not text:
        return text
    text = text.replace("\x00", "")
    # XML 1.0: запрещены большинство C0 control chars кроме tab, CR, LF
    return "".join(
        ch
        for ch in text
        if ch in "\t\n\r" or ord(ch) >= 0x20 or ch in ("\u0085", "\u2028", "\u2029")
    )


def add_folder_row(table, folder_title: str) -> None:
    """Строка-заголовок группы: одна ячейка на всю ширину таблицы."""
    row_cells = table.add_row().cells
    a = row_cells[0]
    b = row_cells[5]
    merged = a.merge(b)
    merged.text = f"Папка «{folder_title}»"
    for p in merged.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)


def add_table_heading_row(table, bold: bool = True) -> None:
    """Строка-шапка с номерами колонок 1…6 (как в методичке)."""
    row = table.add_row().cells
    for i, text in enumerate(["1", "2", "3", "4", "5", "6"], start=0):
        row[i].text = text
        if bold:
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.bold = True


def set_cell_text(cell, text: str, font_size_pt: int = 10) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(font_size_pt)


def build_document(
    source_root: Path,
    output_path: Path,
    *,
    include_listings: bool = True,
    repeat_column_numbers: bool = True,
) -> None:
    files = collect_files(source_root)
    files.sort(key=lambda x: sort_key(x, source_root))

    rows_payload: list[dict] = []
    for fp in files:
        rel = fp.relative_to(source_root)
        rel_posix = rel.as_posix()
        try:
            raw = fp.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = fp.read_text(encoding="utf-8", errors="replace")

        line_count = line_count_for_text(raw)
        size_kb = fp.stat().st_size / 1024
        lang = lang_for_suffix(fp)
        desc = description_for_file(rel_posix, raw, lang)
        rows_payload.append(
            {
                "rel_posix": rel_posix,
                "folder": folder_label(source_root, rel),
                "lines": line_count,
                "size_kb": size_kb,
                "lang": lang,
                "desc": desc,
                "content": raw,
            }
        )

    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run("Таблица 1 – Описание модулей")
    r.bold = True
    r.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"

    hdr = table.add_row().cells
    headers = [
        "№",
        "Название файла",
        "Язык",
        "Строк кода",
        "Вес (КБ)",
        "Описание модуля",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, 10)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True

    add_table_heading_row(table, bold=True)

    current_folder: str | None = None
    n = 0
    for item in rows_payload:
        if item["folder"] != current_folder:
            current_folder = item["folder"]
            add_folder_row(table, current_folder)
            if repeat_column_numbers:
                add_table_heading_row(table, bold=True)

        n += 1
        row = table.add_row().cells
        set_cell_text(row[0], str(n), 9)
        set_cell_text(row[1], item["rel_posix"], 9)
        set_cell_text(row[2], item["lang"], 9)
        set_cell_text(row[3], str(item["lines"]), 9)
        set_cell_text(row[4], format_kb(item["size_kb"]), 9)
        set_cell_text(row[5], item["desc"], 9)

    if include_listings:
        doc.add_page_break()
        h = doc.add_heading("Полный исходный код модулей", level=1)
        h.runs[0].font.size = Pt(14) if h.runs else None

        order_map = {item["rel_posix"]: item["content"] for item in rows_payload}
        for item in rows_payload:
            rel = item["rel_posix"]
            doc.add_heading(rel, level=2)
            p = doc.add_paragraph()
            safe = sanitize_for_word_xml(order_map[rel])
            run = p.add_run(safe)
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="DOCX: таблица модулей + полный код TIP.")
    parser.add_argument(
        "-r",
        "--root",
        type=Path,
        default=DEFAULT_SOURCE_ROOT,
        help="Корень исходников (по умолчанию: ./TIP)",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Имя выходного DOCX",
    )
    parser.add_argument(
        "--no-code",
        action="store_true",
        help="Только таблица, без листингов",
    )
    parser.add_argument(
        "--no-repeat-cols",
        action="store_true",
        help="Не повторять строку 1–6 перед каждой папкой",
    )
    args = parser.parse_args()

    root = args.root.resolve()
    if not root.is_dir():
        raise SystemExit(f"Каталог не найден: {root}")

    build_document(
        root,
        args.output.resolve(),
        include_listings=not args.no_code,
        repeat_column_numbers=not args.no_repeat_cols,
    )
    out = args.output.resolve()
    nfiles = len(collect_files(root))
    print("Готово:", out)
    print("Файлов в таблице:", nfiles)


if __name__ == "__main__":
    main()
