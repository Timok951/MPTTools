"""
Словарь данных БД MPT Tools → DOCX.

Режимы:
  postgres (по умолчанию) — интроспекция реальной СУБД PostgreSQL (каталоги pg_*, COMMENT).
  django — по моделям Django (как раньше).

Таблица: Ключ | Наименование атрибута | Тип данных | Ограничения | Описание

Запуск из корня репозитория:
  .venv\\Scripts\\python.exe generate_data_dictionary.py
  .venv\\Scripts\\python.exe generate_data_dictionary.py --source django
  .venv\\Scripts\\python.exe generate_data_dictionary.py --schema public --skip django_migrations
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
TIP_DIR = ROOT / "TIP"

if str(TIP_DIR) not in sys.path:
    sys.path.insert(0, str(TIP_DIR))

import django  # noqa: E402

os_environ = __import__("os").environ
os_environ.setdefault("DJANGO_SETTINGS_MODULE", "TIP.settings")
django.setup()

from django.apps import apps  # noqa: E402
from django.db import connection  # noqa: E402
from django.db import models  # noqa: E402
from django.utils.encoding import force_str  # noqa: E402

try:
    from docx import Document  # noqa: E402
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
    from docx.shared import Pt  # noqa: E402
except ImportError as e:
    raise SystemExit("Нужен пакет python-docx: pip install python-docx") from e

DEFAULT_APP_LABELS = (
    "core",
    "assets",
    "operations",
    "audit",
    "auth",
    "contenttypes",
    "authtoken",
)


def assert_postgresql() -> None:
    eng = connection.settings_dict.get("ENGINE", "")
    if "postgresql" not in eng:
        raise SystemExit(
            "Режим --source postgres требует подключения к PostgreSQL.\n"
            "В TIP/.env задайте DATABASE_ENGINE=postgresql и параметры БД, либо используйте --source django."
        )


def add_num_row(table, ncols: int = 5) -> None:
    row = table.add_row().cells
    for i in range(ncols):
        cell = row[i]
        cell.text = str(i + 1)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def add_merged_section_row(
    table,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    font_pt: int = 10,
) -> None:
    """Одна строка на всю ширину таблицы (5 колонок) — заголовок раздела."""
    row_cells = table.add_row().cells
    merged = row_cells[0].merge(row_cells[4])
    merged.text = text or ""
    for p in merged.paragraphs:
        for r in p.runs:
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(font_pt)


def append_fk_suffix_for_description(desc: str, fk_line: str | None) -> str:
    """Добавляет строку вида (table.col) под описанием, как в методичных примерах."""
    if not fk_line:
        return desc
    m = re.search(r"FK →\s*([\w.]+)\((\w+)\)", fk_line)
    if not m:
        return desc
    tbl = m.group(1).strip().split(".")[-1]
    col = m.group(2).strip()
    suffix = f"({tbl}.{col})"
    compact = desc.replace("\n", "").replace(" ", "")
    if suffix.replace("(", "").replace(")", "").replace(".", "") in compact.replace(".", ""):
        return desc
    return f"{desc}\n{suffix}"


# --- PostgreSQL introspection ---


def list_base_tables(cursor, schema: str) -> list[str]:
    cursor.execute(
        """
        SELECT c.relname
        FROM pg_catalog.pg_class c
        JOIN pg_catalog.pg_namespace n ON n.oid = c.relnamespace
        WHERE n.nspname = %s
          AND c.relkind = 'r'
        ORDER BY c.relname
        """,
        [schema],
    )
    return [r[0] for r in cursor.fetchall()]


def table_comment(cursor, schema: str, table: str) -> str | None:
    cursor.execute(
        """
        SELECT pg_catalog.obj_description(c.oid, 'pg_class')
        FROM pg_catalog.pg_class c
        JOIN pg_catalog.pg_namespace n ON n.oid = c.relnamespace
        WHERE n.nspname = %s AND c.relname = %s
        """,
        [schema, table],
    )
    row = cursor.fetchone()
    if row and row[0]:
        return str(row[0]).strip()
    return None


def fetch_columns(cursor, schema: str, table: str) -> list[dict]:
    cursor.execute(
        """
        SELECT
            a.attname AS column_name,
            pg_catalog.format_type(a.atttypid, a.atttypmod) AS data_type,
            a.attnotnull AS not_null,
            pg_catalog.pg_get_expr(ad.adbin, ad.adrelid) AS column_default,
            pg_catalog.col_description(a.attrelid, a.attnum) AS col_comment
        FROM pg_catalog.pg_attribute a
        JOIN pg_catalog.pg_class cl ON cl.oid = a.attrelid
        JOIN pg_catalog.pg_namespace ns ON ns.oid = cl.relnamespace
        LEFT JOIN pg_catalog.pg_attrdef ad
            ON ad.adrelid = a.attrelid AND ad.adnum = a.attnum
        WHERE ns.nspname = %s
          AND cl.relname = %s
          AND a.attnum > 0
          AND NOT a.attisdropped
        ORDER BY a.attnum
        """,
        [schema, table],
    )
    cols = []
    for name, dtype, not_null, default, comment in cursor.fetchall():
        cols.append(
            {
                "name": name,
                "data_type": dtype or "",
                "not_null": bool(not_null),
                "default": default,
                "comment": (comment or "").strip() or None,
            }
        )
    return cols


def fetch_pk_columns(cursor, schema: str, table: str) -> set[str]:
    cursor.execute(
        """
        SELECT kcu.column_name
        FROM information_schema.table_constraints tc
        JOIN information_schema.key_column_usage kcu
            ON tc.constraint_schema = kcu.constraint_schema
            AND tc.constraint_name = kcu.constraint_name
            AND tc.table_schema = kcu.table_schema
        WHERE tc.table_schema = %s
          AND tc.table_name = %s
          AND tc.constraint_type = 'PRIMARY KEY'
        ORDER BY kcu.ordinal_position
        """,
        [schema, table],
    )
    return {r[0] for r in cursor.fetchall()}


def _pg_rule(code: str) -> str:
    return {
        "a": "NO ACTION",
        "r": "RESTRICT",
        "c": "CASCADE",
        "n": "SET NULL",
        "d": "SET DEFAULT",
    }.get(code, code)


def fetch_fk_by_column(cursor, schema: str, table: str) -> dict[str, str]:
    """column_name -> текст ограничения FK (целевая таблица/колонка, правила)."""
    cursor.execute(
        """
        SELECT
            att.attname AS column_name,
            fnsp.nspname AS ref_schema,
            frel.relname AS ref_table,
            fatt.attname AS ref_column,
            c.confupdtype::text,
            c.confdeltype::text
        FROM pg_catalog.pg_constraint c
        JOIN pg_catalog.pg_class src ON src.oid = c.conrelid
        JOIN pg_catalog.pg_namespace src_ns ON src_ns.oid = src.relnamespace
        JOIN LATERAL unnest(c.conkey) WITH ORDINALITY AS ck(attnum, ord) ON true
        JOIN pg_catalog.pg_attribute att
            ON att.attrelid = c.conrelid AND att.attnum = ck.attnum
        JOIN LATERAL unnest(c.confkey) WITH ORDINALITY AS fk(attnum, ord2) ON fk.ord2 = ck.ord
        JOIN pg_catalog.pg_attribute fatt
            ON fatt.attrelid = c.confrelid AND fatt.attnum = fk.attnum
        JOIN pg_catalog.pg_class frel ON frel.oid = c.confrelid
        JOIN pg_catalog.pg_namespace fnsp ON fnsp.oid = frel.relnamespace
        WHERE c.contype = 'f'
          AND src_ns.nspname = %s
          AND src.relname = %s
        ORDER BY c.conname, ck.ord
        """,
        [schema, table],
    )
    out: dict[str, str] = {}
    for col, ref_schema, ref_table, ref_col, upd_t, del_t in cursor.fetchall():
        ref = f"{ref_schema}.{ref_table}({ref_col})"
        rule = f"ON UPDATE {_pg_rule(upd_t)} ON DELETE {_pg_rule(del_t)}"
        out[col] = f"FK → {ref}, {rule}"
    return out


def fetch_unique_single_columns(cursor, schema: str, table: str) -> set[str]:
    """Колонки, входящие в уникальное ограничение из одного атрибута (без PK)."""
    cursor.execute(
        """
        SELECT a.attname
        FROM pg_catalog.pg_constraint c
        JOIN pg_catalog.pg_class cl ON cl.oid = c.conrelid
        JOIN pg_catalog.pg_namespace n ON n.oid = cl.relnamespace
        JOIN LATERAL unnest(c.conkey) AS u(attnum) ON true
        JOIN pg_catalog.pg_attribute a ON a.attrelid = c.conrelid AND a.attnum = u.attnum
        WHERE n.nspname = %s
          AND cl.relname = %s
          AND c.contype = 'u'
          AND array_length(c.conkey, 1) = 1
        """,
        [schema, table],
    )
    return {r[0] for r in cursor.fetchall()}


def fetch_table_level_notes(cursor, schema: str, table: str) -> str:
    """CHECK, составные UNIQUE и прочие ограничения (кроме PK/FK — они по колонкам)."""
    cursor.execute(
        """
        SELECT c.conname, c.contype, pg_catalog.pg_get_constraintdef(c.oid, true) AS def
        FROM pg_catalog.pg_constraint c
        JOIN pg_catalog.pg_class cl ON cl.oid = c.conrelid
        JOIN pg_catalog.pg_namespace n ON n.oid = cl.relnamespace
        WHERE n.nspname = %s
          AND cl.relname = %s
          AND c.contype IN ('c', 'u')
          AND NOT (c.contype = 'u' AND coalesce(array_length(c.conkey, 1), 0) = 1)
        ORDER BY c.conname
        """,
        [schema, table],
    )
    parts = []
    for name, ctype, defn in cursor.fetchall():
        tag = "CHECK" if ctype == "c" else "UNIQUE"
        parts.append(f"{tag} {name}: {defn}")
    return "; ".join(parts)


def format_default(val: str | None) -> str | None:
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    if len(s) > 120:
        s = s[:117] + "…"
    return s


# --- Тексты для графы «Описание» (Django + эвристики + COMMENT из БД) ---


def field_description(field: models.Field) -> str:
    bits: list[str] = []
    vn = getattr(field, "verbose_name", None)
    if vn:
        bits.append(force_str(vn))
    ht = getattr(field, "help_text", None)
    if ht:
        bits.append(force_str(ht))
    ch = getattr(field, "choices", None)
    if ch:
        try:
            opts = ", ".join(f"{k}: {force_str(v)}" for k, v in field.flatchoices[:12])
            if len(field.flatchoices) > 12:
                opts += ", …"
            bits.append(f"Допустимые значения: {opts}")
        except Exception:
            pass
    return " ".join(bits) if bits else "—"


def enrich_field_description(field: models.Field) -> str:
    """Расширенное описание поля модели (если в модели не задано — типовые формулировки)."""
    base = field_description(field)
    if base != "—":
        return base
    if getattr(field, "primary_key", False):
        if isinstance(field, (models.AutoField, models.BigAutoField)):
            return "Суррогатный первичный ключ записи (автоинкремент)."
        return "Первичный ключ записи."
    if isinstance(field, (models.ForeignKey, models.OneToOneField)):
        rel = field.remote_field.model
        rlabel = force_str(rel._meta.verbose_name or rel._meta.model_name)
        return f"Связь с сущностью «{rlabel}» (внешний ключ)."
    if isinstance(field, models.BooleanField):
        return "Логический признак (да/нет)."
    if isinstance(field, models.DateTimeField):
        return "Дата и время (с учётом часового пояса)."
    if isinstance(field, models.DateField):
        return "Календарная дата."
    if isinstance(field, models.JSONField):
        return "Структурированные данные в формате JSON."
    if isinstance(field, (models.FileField, models.ImageField)):
        return "Путь к загруженному файлу в хранилище медиа."
    if isinstance(field, models.TextField):
        return "Многострочный текст (без жёсткого лимита длины в БД)."
    if isinstance(field, models.CharField):
        return f"Текстовая строка (до {field.max_length} символов)."
    if isinstance(field, models.IntegerField) and not isinstance(
        field, (models.AutoField, models.BigAutoField)
    ):
        return "Целое число."
    if isinstance(field, models.DecimalField):
        return (
            f"Десятичное число (точность {field.max_digits}, дробная часть {field.decimal_places})."
        )
    if isinstance(field, models.EmailField):
        return "Адрес электронной почты."
    return "—"


# Доп. подписи таблиц и ячеек (редкие случаи, когда в модели мало текста).
EXTRA_TABLE_LABELS: dict[str, str] = {
    "django_migrations": "Журнал применённых миграций схемы БД (Django)",
    "django_session": "Серверные сессии пользователей веб-приложения",
    "django_admin_log": "Журнал действий в административной панели Django",
    "django_content_type": "Справочник типов объектов (приложение + модель) для GenericForeignKey",
}

EXTRA_COLUMN_HELP: dict[tuple[str, str], str] = {
    ("django_migrations", "id"): "Первичный ключ записи о миграции.",
    ("django_migrations", "app"): "Имя приложения Django, к которому относится миграция.",
    ("django_migrations", "name"): "Имя файла миграции (класс миграции).",
    ("django_migrations", "applied"): "Дата и время применения миграции к этой базе.",
    ("django_session", "session_key"): "Уникальный ключ сессии в cookie клиента.",
    ("django_session", "session_data"): "Сериализованные данные сессии.",
    ("django_session", "expire_date"): "Момент истечения сессии.",
    ("auth_user", "password"): "Хеш пароля пользователя (формат Django: алгоритм$итерации$соль$хеш).",
    ("auth_user", "is_superuser"): "Признак суперпользователя (полный доступ к админке и правам).",
    ("auth_user", "username"): "Уникальное имя учётной записи для входа.",
    ("auth_user", "is_staff"): "Допуск к входу в административную панель.",
    ("auth_user", "is_active"): "Учётная запись может использоваться для аутентификации.",
    ("auth_user", "date_joined"): "Дата и время регистрации пользователя в системе.",
    ("auth_user", "last_login"): "Время последнего успешного входа.",
    ("auth_user", "first_name"): "Имя (необязательно).",
    ("auth_user", "last_name"): "Фамилия (необязательно).",
    ("auth_user", "email"): "Адрес электронной почты.",
    ("authtoken_token", "key"): "Секретный ключ токена API (передаётся клиентом в заголовках).",
    ("authtoken_token", "created"): "Когда был выдан токен.",
    ("authtoken_token", "user_id"): "Пользователь-владелец токена REST API.",
}


KNOWN_COLUMN_HINTS: dict[str, str] = {
    "created_at": "Дата и время создания записи.",
    "updated_at": "Дата и время последнего изменения записи.",
    "deleted_at": "Мягкое удаление: если заполнено — запись скрыта из обычных выборок.",
    "requested_at": "Когда была создана заявка.",
    "processed_at": "Когда заявка была обработана (одобрена/отклонена).",
    "used_at": "Момент операции списания/использования.",
    "read_at": "Когда сообщение было прочитано получателем.",
    "last_login": "Последний успешный вход в систему.",
    "date_joined": "Дата регистрации пользователя.",
    "object_id": "Идентификатор объекта в таблице целевой модели (строка для универсальности).",
    "object_repr": "Краткое текстовое представление объекта на момент записи в журнал.",
    "meta": "Дополнительные структурированные данные события (JSON).",
}


def heuristic_column_description(
    column_name: str,
    *,
    is_pk: bool,
    in_fk: bool,
    data_type: str,
) -> str | None:
    if is_pk and column_name == "id":
        return "Первичный ключ (числовой идентификатор строки)."
    if column_name in KNOWN_COLUMN_HINTS:
        return KNOWN_COLUMN_HINTS[column_name]
    if column_name.endswith("_id") and in_fk:
        return (
            "Внешний ключ: ссылка на запись в связанной таблице "
            "(см. также графу «Ограничения»)."
        )
    if column_name in ("slug", "entity_slug"):
        return "Краткий программный идентификатор сущности в портале."
    if "json" in data_type.lower() or data_type.upper() == "JSONB":
        return "Данные в формате JSON (структура определяется прикладной логикой)."
    return None


def merge_description_cells(
    django_text: str | None,
    pg_comment: str | None,
    heuristic: str | None,
) -> str:
    parts: list[str] = []
    if django_text and django_text != "—":
        parts.append(django_text)
    if pg_comment:
        parts.append(f"[из БД] {pg_comment}")
    if not parts and heuristic:
        parts.append(heuristic)
    if not parts:
        return "—"
    return " ".join(parts)


def collect_django_description_maps() -> tuple[dict[str, str], dict[tuple[str, str], str]]:
    """Соответствие имён таблиц/колонок БД текстам из моделей Django."""
    table_labels: dict[str, str] = {}
    column_text: dict[tuple[str, str], str] = {}
    for model in apps.get_models():
        if model._meta.proxy:
            continue
        db_table = model._meta.db_table
        label = force_str(model._meta.verbose_name_plural or model._meta.verbose_name or db_table)
        table_labels[db_table] = label
        for field in model._meta.get_fields():
            col = getattr(field, "column", None)
            if not col or field.many_to_many or field.one_to_many:
                continue
            text = enrich_field_description(field)
            if text and text != "—":
                column_text[(db_table, col)] = text
    table_labels.update(EXTRA_TABLE_LABELS)
    column_text.update(EXTRA_COLUMN_HELP)
    return table_labels, column_text


def build_doc_postgres(
    output: Path,
    *,
    schemas: tuple[str, ...],
    skip_tables: frozenset[str],
) -> int:
    assert_postgresql()
    django_table_labels, django_column_text = collect_django_description_maps()
    headers = ("Ключ", "Наименование атрибута", "Тип данных", "Ограничения", "Описание")
    doc = Document()
    h = doc.add_heading("Словарь данных базы данных (MPT Tools)", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Таблицы и атрибуты получены интроспекцией PostgreSQL "
        "(pg_catalog, information_schema). Типы — как возвращает format_type(). "
        "Графа «Описание» заполняется по метаданным моделей Django (verbose_name, help_text, варианты choices), "
        "дополняется комментариями COMMENT из БД (если есть) и краткими пояснениями для типовых имён колонок. "
        "Все таблицы перечислены в одной таблице документа; каждая сущность начинается с объединённых строк "
        "(имя таблицы и пояснение)."
    )

    main_tbl = doc.add_table(rows=0, cols=5)
    main_tbl.style = "Table Grid"
    hr = main_tbl.add_row().cells
    for i, htxt in enumerate(headers):
        hr[i].text = htxt
        for p in hr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    add_num_row(main_tbl, 5)

    total_tables = 0
    with connection.cursor() as cursor:
        for schema in schemas:
            tables = [t for t in list_base_tables(cursor, schema) if t not in skip_tables]
            for table in tables:
                total_tables += 1
                t_comment = table_comment(cursor, schema, table)
                add_merged_section_row(
                    main_tbl,
                    f"{schema}.{table}",
                    bold=True,
                    italic=False,
                    font_pt=11,
                )
                sub = (
                    django_table_labels.get(table)
                    or t_comment
                    or f"Таблица «{table}», схема «{schema}»."
                )
                add_merged_section_row(main_tbl, sub, bold=False, italic=True, font_pt=10)

                pk = fetch_pk_columns(cursor, schema, table)
                fk_map = fetch_fk_by_column(cursor, schema, table)
                uniq = fetch_unique_single_columns(cursor, schema, table)
                cols = fetch_columns(cursor, schema, table)
                table_notes = fetch_table_level_notes(cursor, schema, table)

                for col in cols:
                    name = col["name"]
                    key_parts: list[str] = []
                    if name in pk:
                        key_parts.append("PK")
                    if name in fk_map:
                        key_parts.append("FK")
                    key = ", ".join(key_parts) if key_parts else ""

                    constr_bits: list[str] = []
                    if col["not_null"]:
                        constr_bits.append("NOT NULL")
                    if name in uniq and name not in pk:
                        constr_bits.append("UNIQUE")
                    d = format_default(col["default"])
                    if d:
                        constr_bits.append(f"DEFAULT {d}")
                    if name in fk_map:
                        constr_bits.append(fk_map[name])

                    dj = django_column_text.get((table, name))
                    heur = heuristic_column_description(
                        name,
                        is_pk=name in pk,
                        in_fk=name in fk_map,
                        data_type=col["data_type"],
                    )
                    desc_cell = merge_description_cells(dj, col["comment"], heur)
                    desc_cell = append_fk_suffix_for_description(desc_cell, fk_map.get(name))

                    row = main_tbl.add_row().cells
                    row[0].text = key
                    row[1].text = name
                    row[2].text = col["data_type"]
                    row[3].text = ", ".join(constr_bits)
                    row[4].text = desc_cell
                    for c in row:
                        for p in c.paragraphs:
                            for rr in p.runs:
                                rr.font.size = Pt(8)

                if table_notes:
                    add_merged_section_row(
                        main_tbl,
                        f"Ограничения уровня таблицы: {table_notes}",
                        bold=False,
                        italic=True,
                        font_pt=9,
                    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return total_tables


# --- Django (модели) — прежняя логика ---


def pg_type_for_field(field: models.Field) -> str:
    if isinstance(field, models.BigAutoField):
        return "BigSerial / BIGINT"
    if isinstance(field, models.AutoField):
        return "Serial / INTEGER"
    if isinstance(field, models.BigIntegerField):
        return "BIGINT"
    if isinstance(field, models.SmallIntegerField):
        return "SMALLINT"
    if isinstance(field, models.IntegerField):
        return "INTEGER"
    if isinstance(field, models.PositiveIntegerField):
        return "INTEGER"
    if isinstance(field, models.PositiveSmallIntegerField):
        return "SMALLINT"
    if isinstance(field, models.BooleanField):
        return "BOOLEAN"
    if isinstance(field, models.DateTimeField):
        return "TIMESTAMPTZ"
    if isinstance(field, models.DateField):
        return "DATE"
    if isinstance(field, models.TimeField):
        return "TIME"
    if isinstance(field, models.DecimalField):
        return f"NUMERIC({field.max_digits},{field.decimal_places})"
    if isinstance(field, models.FloatField):
        return "DOUBLE PRECISION"
    if isinstance(field, models.EmailField):
        return f"VARCHAR({field.max_length})"
    if isinstance(field, models.URLField):
        return f"VARCHAR({field.max_length})"
    if isinstance(field, models.UUIDField):
        return "UUID"
    if isinstance(field, models.JSONField):
        return "JSONB"
    if isinstance(field, models.TextField):
        return "TEXT"
    if isinstance(field, models.CharField):
        return f"VARCHAR({field.max_length})"
    if isinstance(field, models.BinaryField):
        return "BYTEA"
    if isinstance(field, models.FileField) or isinstance(field, models.ImageField):
        return f"VARCHAR({field.max_length})"
    if isinstance(field, models.ForeignKey):
        return pg_type_for_field(field.target_field)
    if isinstance(field, models.OneToOneField):
        return pg_type_for_field(field.target_field)
    return field.get_internal_type()


def on_delete_code(field: models.ForeignKey | models.OneToOneField) -> str:
    return field.remote_field.on_delete.__name__.replace("_", " ")


def constraint_parts_for_field(field: models.Field, model: type[models.Model]) -> list[str]:
    parts: list[str] = []
    if not getattr(field, "null", False):
        parts.append("NOT NULL")
    if getattr(field, "unique", False) and not getattr(field, "primary_key", False):
        parts.append("UNIQUE")
    if getattr(field, "db_index", False) and not getattr(field, "primary_key", False):
        parts.append("INDEX")

    if field.has_default():
        d = field.default
        if d is models.NOT_PROVIDED:
            pass
        elif callable(d):
            name = getattr(d, "__name__", "callable")
            parts.append(f"DEFAULT ({name}())")
        elif isinstance(d, bool):
            parts.append(f"DEFAULT {'TRUE' if d else 'FALSE'}")
        elif d is None:
            parts.append("DEFAULT NULL")
        else:
            parts.append(f"DEFAULT {repr(d)}")

    if isinstance(field, (models.ForeignKey, models.OneToOneField)):
        rel_table = field.remote_field.model._meta.db_table
        rel_col = field.target_field.column
        parts.append(f"FK → {rel_table}({rel_col})")
        parts.append(f"ON DELETE {on_delete_code(field)}")

    return parts


def key_for_field(field: models.Field) -> str:
    if getattr(field, "primary_key", False):
        return "PK"
    if isinstance(field, (models.ForeignKey, models.OneToOneField)):
        return "FK"
    return ""


def table_level_constraints(model: type[models.Model]) -> str:
    lines: list[str] = []
    ut = getattr(model._meta, "unique_together", None)
    if ut:
        for combo in ut:
            lines.append(f"UNIQUE ({', '.join(combo)})")
    for c in getattr(model._meta, "constraints", []) or []:
        lines.append(str(c))
    return "; ".join(lines) if lines else ""


def collect_models(app_labels: tuple[str, ...], *, include_m2m_through: bool = True) -> list[type[models.Model]]:
    seen: set[type[models.Model]] = set()
    out: list[type[models.Model]] = []
    for label in app_labels:
        try:
            app_config = apps.get_app_config(label)
        except LookupError:
            continue
        for model in app_config.get_models():
            if model in seen:
                continue
            if model._meta.proxy:
                continue
            seen.add(model)
            out.append(model)

    if include_m2m_through:
        for model in list(out):
            for m2m in model._meta.local_many_to_many:
                through = m2m.remote_field.through
                if through and through not in seen and not through._meta.proxy:
                    seen.add(through)
                    out.append(through)

    out.sort(key=lambda m: (m._meta.app_label, m._meta.db_table))
    return out


def build_doc_django(models_list: list[type[models.Model]], output: Path) -> None:
    doc = Document()
    h = doc.add_heading("Словарь данных базы данных (MPT Tools)", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Таблицы и атрибуты сформированы по моделям Django (целевая СУБД — PostgreSQL). "
        "Типы приведены в нотации, близкой к PostgreSQL. Все таблицы — в одном листинге; "
        "разделы выделены объединёнными строками."
    )
    headers = ("Ключ", "Наименование атрибута", "Тип данных", "Ограничения", "Описание")

    main_tbl = doc.add_table(rows=0, cols=5)
    main_tbl.style = "Table Grid"
    hr = main_tbl.add_row().cells
    for i, htxt in enumerate(headers):
        hr[i].text = htxt
        for p in hr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    add_num_row(main_tbl, 5)

    for model in models_list:
        meta = model._meta
        title_ru = force_str(meta.verbose_name_plural or meta.verbose_name or meta.model)
        add_merged_section_row(main_tbl, meta.db_table, bold=True, italic=False, font_pt=11)
        add_merged_section_row(main_tbl, title_ru, bold=False, italic=True, font_pt=10)

        for field in meta.get_fields():
            if not getattr(field, "column", None):
                continue
            if field.many_to_many or field.one_to_many:
                continue
            col_name = field.column
            key = key_for_field(field)
            sql_t = pg_type_for_field(field)
            constr = ", ".join(constraint_parts_for_field(field, model))
            desc = enrich_field_description(field)
            if desc == "—":
                desc = field_description(field)
            if isinstance(field, (models.ForeignKey, models.OneToOneField)):
                rel_t = field.remote_field.model._meta.db_table
                rel_c = field.target_field.column
                desc = append_fk_suffix_for_description(
                    desc, f"FK → {rel_t}({rel_c})"
                )
            row = main_tbl.add_row().cells
            row[0].text = key
            row[1].text = col_name
            row[2].text = sql_t
            row[3].text = constr
            row[4].text = desc
            for c in row:
                for p in c.paragraphs:
                    for rr in p.runs:
                        rr.font.size = Pt(8)

        extra = table_level_constraints(model)
        if extra:
            add_merged_section_row(
                main_tbl,
                f"Ограничения уровня таблицы: {extra}",
                bold=False,
                italic=True,
                font_pt=9,
            )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Словарь данных → DOCX: PostgreSQL (интроспекция) или Django-модели."
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=ROOT / "Словарь_данных_MPT_Tools.docx",
        help="Путь к выходному DOCX",
    )
    parser.add_argument(
        "--source",
        choices=("postgres", "django"),
        default="postgres",
        help="Источник метаданных (по умолчанию: postgres)",
    )
    parser.add_argument(
        "--schema",
        type=str,
        default="public",
        help="Схема PostgreSQL (можно несколько через запятую)",
    )
    parser.add_argument(
        "--skip",
        type=str,
        default="",
        help="Имена таблиц через запятую, исключить из отчёта (только postgres)",
    )
    parser.add_argument(
        "--apps",
        type=str,
        default=",".join(DEFAULT_APP_LABELS),
        help="--source django: app_label через запятую",
    )
    parser.add_argument(
        "--no-m2m-through",
        action="store_true",
        help="--source django: не добавлять связующие M2M",
    )
    args = parser.parse_args()
    out = args.output.resolve()

    if args.source == "postgres":
        schemas = tuple(s.strip() for s in args.schema.split(",") if s.strip())
        if not schemas:
            raise SystemExit("Укажите хотя бы одну схему в --schema")
        skip = frozenset(s.strip() for s in args.skip.split(",") if s.strip())
        n = build_doc_postgres(out, schemas=schemas, skip_tables=skip)
        print("Готово:", out)
        print("Таблиц в отчёте:", n)
        return

    labels = tuple(s.strip() for s in args.apps.split(",") if s.strip())
    models_list = collect_models(labels, include_m2m_through=not args.no_m2m_through)
    if not models_list:
        raise SystemExit("Не найдено ни одной модели. Проверьте --apps.")
    build_doc_django(models_list, out)
    print("Готово:", out)
    print("Таблиц (моделей):", len(models_list))


if __name__ == "__main__":
    main()
