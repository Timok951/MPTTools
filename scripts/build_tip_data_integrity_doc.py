# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_cell_border(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_row(table, row_data):
    cells = table.add_row().cells
    for i, text in enumerate(row_data):
        p = cells[i].paragraphs[0]
        p.add_run(text)
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    template = root / "Контроль целостности данных.docx"
    doc = Document(str(template)) if template.exists() else Document()
    table_style = doc.tables[0].style if doc.tables else None

    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    title = doc.add_paragraph("Контроль целостности данных для проекта TIP")
    title.runs[0].bold = True
    title.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=6)
    if table_style is not None:
        table.style = table_style
    table.autofit = False
    table.columns[0].width = Cm(0.9)
    table.columns[1].width = Cm(3.9)
    table.columns[2].width = Cm(4.4)
    table.columns[3].width = Cm(4.8)
    table.columns[4].width = Cm(3.8)
    table.columns[5].width = Cm(3.8)

    headers = [
        "№",
        "Этап валидации / Объект",
        "Тип проверки / ограничения",
        "Описание",
        "Действие при ошибке / нарушении",
        "Класс / Страница / Файл",
    ]
    for i, text in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ("1", "Регистрация: email", "Формат email + разрешенный домен", "Проверка корректности почты и принадлежности домену из whitelist.", "Отказ в регистрации, текст ошибки в форме.", "core.models.RegistrationAllowedEmailDomain / register.html / inventory/forms.py"),
        ("2", "Регистрация: username", "Regex + длина + уникальность", "Логин проверяется на допустимые символы и отсутствие дубликатов.", "Поле помечается ошибкой, запись не создается.", "inventory/forms.py (RussianUserCreationForm)"),
        ("3", "Регистрация: пароль", "Сложность и подтверждение", "Сравнение пароля и подтверждения, применение валидаторов Django.", "Ошибка формы, пользователь не создается.", "inventory/forms.py / auth views"),
        ("4", "Кабинет: name", "Только цифры", "Название кабинета допускает только числовые значения.", "ValidationError, блокировка сохранения.", "core.models.Cabinet.clean"),
        ("5", "Кабинет: floor", "Пусто или только цифры", "Этаж в строковом поле проверяется как числовой код.", "ValidationError по полю floor.", "core.models.Cabinet.clean"),
        ("6", "Категория оборудования", "name unique", "Наименование категории не должно дублироваться.", "Ошибка БД/формы, отмена операции.", "core.models.EquipmentCategory"),
        ("7", "Оборудование: inventory_number", "Уникальность", "Каждая карточка имеет уникальный инвентарный номер.", "IntegrityError/ошибка формы.", "assets.models.Equipment"),
        ("8", "Оборудование: serial_number", "Regex [буквы/цифры/-_/]", "Проверка допустимых символов серийного номера.", "ValidationError serial_number.", "assets.models.Equipment.clean"),
        ("9", "Оборудование: model", "Regex + длина", "Модель допускает буквы, цифры, пробел и -_/.", "ValidationError model.", "assets.models.Equipment.clean"),
        ("10", "Оборудование: quantity_available", "<= quantity_total", "Доступный остаток не может превышать общий.", "ValidationError, сохранение отменяется.", "assets.models.Equipment.clean"),
        ("11", "Оборудование: purchase_date", "Не в будущем, разумный минимум", "Исключаются нереалистичные и будущие даты покупки.", "ValidationError purchase_date.", "assets.models.Equipment.clean"),
        ("12", "Оборудование: warranty_end", ">= purchase_date, не слишком далеко", "Контроль логики гарантийной даты.", "ValidationError warranty_end.", "assets.models.Equipment.clean"),
        ("13", "Оборудование расходник", "quantity_available = quantity_total", "Для расходников остаток синхронизируется с общим количеством.", "Автокорректировка перед сохранением.", "assets/signals.py pre_save"),
        ("14", "Заявка: quantity", ">0, спец-исключение quantity=0", "0 допустим только для set_in_stock у нерасходника при пополнении.", "ValidationError quantity.", "operations.models.EquipmentRequest.clean"),
        ("15", "Заявка: needed_by", "Не раньше сегодня, не дальше +365", "Контроль диапазона плановой даты.", "ValidationError needed_by.", "operations.models.EquipmentRequest.clean"),
        ("16", "Заявка: comment", "Осмысленный текст, <=500", "Запрещены пустые/символьные комментарии без букв/цифр.", "ValidationError comment.", "operations.models.EquipmentRequest.clean"),
        ("17", "Заявка: request_kind", "Только restock/writeoff", "Ограничение выбора типа заявки по enum.", "Ошибка формы/API, отклонение payload.", "operations/models.py + serializers"),
        ("18", "Заявка: status", "pending/approved/rejected", "Только поддерживаемые статусы жизненного цикла.", "Ошибка валидации статуса.", "operations.models.EquipmentRequest"),
        ("19", "Сообщение в заявке", "request + author обязательны", "Сообщение всегда связано с заявкой и автором.", "Ошибка сохранения/валидации.", "operations.models.EquipmentRequestMessage"),
        ("20", "Отметка прочтения треда", "unique(user, request)", "Одна запись отметки на пользователя и заявку.", "IntegrityError при дубле.", "operations.models.EquipmentRequestThreadRead"),
        ("21", "Фото заявки", "image обязательное, caption <=200", "Проверка наличия файла и ограничение подписи.", "Ошибка формы/serializer.", "operations.models.EquipmentRequestPhoto"),
        ("22", "Личное сообщение", "sender != recipient", "Запрет отправки сообщения самому себе.", "CheckConstraint violation.", "core.models.DirectMessage"),
        ("23", "MaterialUsage.quantity", ">0", "Списание расходников только положительным количеством.", "ValidationError quantity.", "operations.models.MaterialUsage.clean"),
        ("24", "PeriodicSchedule.quantity", ">0", "Периодическая операция только с положительным объемом.", "ValidationError quantity.", "operations.models.PeriodicMaterialUsageSchedule.clean"),
        ("25", "PeriodicSchedule.equipment", "Только расходник", "Расписание запрещено для нерасходуемых позиций.", "ValidationError equipment.", "operations.models.PeriodicMaterialUsageSchedule.clean"),
        ("26", "InventoryAdjustment.delta", "Итоговые остатки не <0", "Корректировка не должна уводить total/available в минус.", "ValidationError, rollback транзакции.", "assets.models.InventoryAdjustment.clean"),
        ("27", "EquipmentCheckout.related_request", "Только approved заявка", "Выдача допустима только по одобренной заявке.", "ValidationError related_request.", "assets.models.EquipmentCheckout.clean"),
        ("28", "EquipmentCheckout.quantity", "<= request.quantity", "Количество выдачи не больше количества в заявке.", "ValidationError quantity.", "assets.models.EquipmentCheckout.clean"),
        ("29", "EquipmentCheckout.returned_at", ">= taken_at", "Дата возврата не раньше даты выдачи.", "ValidationError returned_at.", "assets.models.EquipmentCheckout.clean"),
        ("30", "Права API", "Роль + объектный доступ", "Проверка role_matrix и owner/privileged политики.", "HTTP 403 Forbidden.", "inventory/api/permissions.py, viewsets.py"),
        ("31", "Валидация API create/update", "instance.full_clean()", "Модельные правила принудительно применяются в serializer.", "HTTP 400 с деталями полей.", "inventory/api/serializers.py AuditActorModelSerializer"),
        ("32", "Частичный PATCH API", "exclude непереданных полей", "Для partial update исключаются неприсланные поля из full_clean.", "HTTP 400 по переданным некорректным полям.", "inventory/api/serializers.py"),
        ("33", "Роли UI", "user_has_capability/user_in_group", "Ограничение разделов и действий по роли.", "403/redirect + warning message.", "inventory/views.py, inventory/authz.py"),
        ("34", "Нерабочее время", "Schedule guard", "Операции блокируются в нерабочий период согласно графику.", "UI-предупреждение, действие не выполняется.", "inventory/middleware.py"),
        ("35", "Импорт backup", "Проверка формата/результата", "Контроль корректности входного файла и этапов восстановления.", "Диагностическое сообщение + лог.", "inventory/views.py backup/import handlers"),
        ("36", "Аудит изменений", "Фиксация old/new + actor", "Каждое изменение сущности журналируется с метаданными.", "Запись в аудит даже при soft delete.", "audit/models.py + signals"),
        ("37", "Soft delete", "deleted_at timestamp", "Физическое удаление заменено логическим для ряда сущностей.", "Запись скрывается из active manager.", "core.models.SoftDeleteModel"),
        ("38", "Email уведомления", "Проверка email адресатов", "Письма отправляются только валидным адресатам, ошибки не роняют процесс.", "Логирование ошибки, пропуск получателя.", "core/message_email_notify.py"),
        ("39", "Формы портала", "Field + clean() + non_field_errors", "Проверки для CRUD-портала: текст, даты, числовые диапазоны.", "Ошибка формы с явным сообщением.", "inventory/portal_forms.py, portal/object_form.html"),
        ("40", "Отчеты/экспорт", "Валидный период и параметры", "Перед выгрузкой проверяются даты, фильтры и доступ роли.", "Предупреждение в UI/API, выгрузка не формируется.", "inventory/views.py (analytics, export, print)"),
    ]
    for row in rows:
        add_row(table, row)

    for tr in table.rows:
        for tc in tr.cells:
            set_cell_border(tc)
            for p in tc.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1
                for run in p.runs:
                    run.font.size = Pt(9.5)

    out_path = root / "TIP_Контроль_целостности_данных.docx"
    doc.save(str(out_path))
    print(out_path.resolve())


if __name__ == "__main__":
    main()
